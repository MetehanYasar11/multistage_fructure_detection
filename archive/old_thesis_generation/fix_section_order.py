"""
Fix section order: Current is 1-3-2, should be 1-2-3
Reorder the sections properly in the DOCX
"""

from docx import Document
from pathlib import Path

def fix_section_order():
    """Reorder sections from 1-3-2 to 1-2-3"""
    
    # Load the incorrectly ordered document
    doc = Document('THESIS_SECTIONS_1_2_3_COMPLETE.docx')
    
    # Create new document with correct order
    new_doc = Document()
    
    # Copy styles
    for style in doc.styles:
        if style.type == 1:  # Paragraph style
            try:
                new_doc.styles.add_style(style.name, style.type)
            except:
                pass
    
    # Track sections
    section1_end = -1
    section2_start = -1
    section2_end = -1
    section3_start = -1
    
    # Find section boundaries
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading 1'):
            heading_text = para.text.strip()
            
            if heading_text.startswith('1.'):
                print(f"Found Section 1 at para {i}: {heading_text}")
                
            elif heading_text.startswith('3.'):
                print(f"Found Section 3 at para {i}: {heading_text}")
                section1_end = i - 1
                section3_start = i
                
            elif heading_text.startswith('2.'):
                print(f"Found Section 2 at para {i}: {heading_text}")
                section2_start = i
    
    # If Section 2 comes after Section 3, we need to reorder
    if section2_start > section3_start and section3_start > 0:
        print(f"\nReordering needed:")
        print(f"  Section 1: paras 0-{section1_end}")
        print(f"  Section 3: paras {section3_start}-{section2_start-1}")
        print(f"  Section 2: paras {section2_start}-{len(doc.paragraphs)-1}")
        
        # Find where Section 3 ends (before Section 2)
        section3_end = section2_start - 1
        
        # Copy in correct order: 1, then 2, then 3
        
        # Copy Section 1 (0 to section1_end)
        print(f"\nCopying Section 1 ({section1_end + 1} paragraphs)...")
        for i in range(0, section1_end + 1):
            copy_paragraph(doc.paragraphs[i], new_doc)
        
        # Copy Section 2 (section2_start to end)
        print(f"Copying Section 2 ({len(doc.paragraphs) - section2_start} paragraphs)...")
        for i in range(section2_start, len(doc.paragraphs)):
            copy_paragraph(doc.paragraphs[i], new_doc)
        
        # Copy Section 3 (section3_start to section3_end)
        print(f"Copying Section 3 ({section3_end - section3_start + 1} paragraphs)...")
        for i in range(section3_start, section3_end + 1):
            copy_paragraph(doc.paragraphs[i], new_doc)
        
        # Save corrected document
        output_path = Path('THESIS_SECTIONS_1_2_3_CORRECTED_ORDER.docx')
        new_doc.save(output_path)
        print(f"\nCorrected document saved to: {output_path.absolute()}")
        print(f"Total paragraphs: {len(new_doc.paragraphs)}")
        print(f"Order: Section 1 -> Section 2 -> Section 3")
        
        return new_doc, output_path
    else:
        print("\nOrder is already correct (1-2-3). No changes needed.")
        return doc, Path('THESIS_SECTIONS_1_2_3_COMPLETE.docx')

def copy_paragraph(source_para, target_doc):
    """Copy paragraph with formatting"""
    
    # Handle page breaks
    if '\f' in source_para.text:
        target_doc.add_page_break()
        return
    
    # Create paragraph with same style
    if source_para.style.name.startswith('Heading'):
        new_para = target_doc.add_heading(level=int(source_para.style.name[-1]))
        new_para.text = source_para.text
    elif source_para.style.name == 'List Bullet':
        new_para = target_doc.add_paragraph(source_para.text, style='List Bullet')
    else:
        new_para = target_doc.add_paragraph()
        
        # Copy runs with formatting
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
    
    # Copy alignment
    new_para.alignment = source_para.alignment

if __name__ == "__main__":
    doc, path = fix_section_order()
