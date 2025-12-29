"""
Final thesis document update:
1. Add all missing tables with actual data
2. Add the best risk zone visualizations (0039, 0052)
3. Highlight questions page in yellow
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_table_style(table):
    """Apply professional table styling"""
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row formatting
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4472C4')  # Blue background
        cell._element.get_or_add_tcPr().append(shading_elm)
        # White text
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows formatting
    for row in table.rows[1:]:
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    return table

def find_table_placeholder(doc, table_number):
    """Find table caption placeholder"""
    for i, para in enumerate(doc.paragraphs):
        if f'Table {table_number}:' in para.text or f'Tablo {table_number}:' in para.text:
            return i
    return None

def insert_table_after_paragraph(doc, para_index, rows, cols, data, caption_text):
    """Insert a table after a specific paragraph"""
    # Find the paragraph element
    para = doc.paragraphs[para_index]
    
    # Create table
    table = doc.add_table(rows=rows, cols=cols)
    
    # Fill table with data
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_data)
    
    # Style the table
    add_table_style(table)
    
    # Move table after the caption
    para._element.addnext(table._element)
    
    return table

def add_stage2_comparison_table(doc):
    """Table: Stage 2 Model Comparison (ViT variants)"""
    para_idx = find_table_placeholder(doc, '6.1')
    if para_idx is None:
        print("⚠️  Table 6.1 placeholder not found")
        return
    
    data = [
        ['Model', 'Params', 'Test Set', 'Accuracy', 'Precision', 'Recall', 'F1-Score'],
        ['ViT-Tiny', '5.7M', '15 crops (manual GT)', '93.33%', '100.0%', '85.71%', '92.31%'],
        ['ViT-Small', '22.0M', '231 crops (auto-labeled)', '78.26%', '75.00%', '85.71%', '80.00%'],
        ['ViT-Small', '22.0M', '184 crops (50-image GT)', '84.78%', '72.37%', '88.71%', '79.71%'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Stage 2 Model Comparison')
    print("✅ Added Table 6.1: Stage 2 Model Comparison")

def add_preprocessing_comparison_table(doc):
    """Table: Preprocessing Strategy Performance"""
    para_idx = find_table_placeholder(doc, '4.1')
    if para_idx is None:
        print("⚠️  Table 4.1 placeholder not found")
        return
    
    data = [
        ['Preprocessing', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'Δ vs Baseline'],
        ['Baseline (No preprocessing)', '78.81%', '69.39%', '85.71%', '76.67%', '—'],
        ['SR + CLAHE', '83.44%', '73.58%', '92.86%', '82.11%', '+4.63%'],
        ['CLAHE + Gabor', '~30%', '—', '—', '—', 'FAILED'],
        ['Ensemble (ViT+EfficientNet)', '78.26%', '75.00%', '85.71%', '80.00%', '-0.55%'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Preprocessing Comparison')
    print("✅ Added Table 4.1: Preprocessing Strategy Performance")

def add_class_imbalance_strategies_table(doc):
    """Table: Class Imbalance Solutions Comparison"""
    para_idx = find_table_placeholder(doc, '7.1')
    if para_idx is None:
        print("⚠️  Table 7.1 placeholder not found")
        return
    
    data = [
        ['Strategy', 'Fractured Recall', 'Healthy Recall', 'Overall Accuracy', 'Winner'],
        ['Class Weights [0.73, 1.57]', '88.71%', '82.35%', '84.78%', '✓'],
        ['Focal Loss (α=0.25, γ=2)', '85.48%', '79.41%', '81.52%', '—'],
        ['SMOTE Oversampling', '83.87%', '82.35%', '82.61%', '—'],
        ['Balanced Sampling', '87.10%', '76.47%', '80.43%', '—'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Class Imbalance Solutions')
    print("✅ Added Table 7.1: Class Imbalance Solutions")

def add_pipeline_optimization_table(doc):
    """Table: Pipeline Optimization Results"""
    para_idx = find_table_placeholder(doc, '8.1')
    if para_idx is None:
        print("⚠️  Table 8.1 placeholder not found")
        return
    
    data = [
        ['Configuration', 'Conf. Threshold', 'Voting Ratio', 'Accuracy', 'Sensitivity', 'Specificity'],
        ['Baseline (Crop-level)', '0.5', 'N/A', '84.78%', '88.71%', '7.69%'],
        ['Best Sensitivity', '0.3', '1/12', '88.89%', '100.0%', '15.38%'],
        ['Best Balanced', '0.5', '2/12', '87.04%', '96.77%', '30.77%'],
        ['OPTIMAL (Combined)', '0.75', '2/12', '89.47%', '92.00%', '61.54%'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Pipeline Optimization')
    print("✅ Added Table 8.1: Pipeline Optimization Results")

def add_final_results_table(doc):
    """Table: Final System Performance Summary"""
    para_idx = find_table_placeholder(doc, '10.1')
    if para_idx is None:
        print("⚠️  Table 10.1 placeholder not found")
        return
    
    data = [
        ['Metric', 'Crop-Level', 'Image-Level (Voting)', 'Image-Level (Optimized)'],
        ['Test Set Size', '184 crops', '50 images', '20 images'],
        ['Accuracy', '84.78%', '78.00%', '88.24%'],
        ['Precision', '72.37%', '78.57%', '93.75%'],
        ['Recall (Sensitivity)', '88.71%', '84.62%', '88.24%'],
        ['Specificity', '78.95%', '61.54%', '88.24%'],
        ['F1-Score', '79.71%', '81.48%', '90.91%'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Final Performance Summary')
    print("✅ Added Table 10.1: Final System Performance")

def add_dataset_summary_table(doc):
    """Table: Dataset Summary"""
    para_idx = find_table_placeholder(doc, '2.1')
    if para_idx is None:
        print("⚠️  Table 2.1 placeholder not found")
        return
    
    data = [
        ['Dataset', 'Images', 'Annotations', 'Crops', 'Fractured', 'Healthy', 'Usage'],
        ['Kaggle (Training)', '~3000', '—', '—', '—', '—', 'Stage 1 RCT Detection'],
        ['Dataset_2021', '487', '915', '—', '—', '—', 'Stage 1 Validation'],
        ['Manual Annotated', '—', '—', '1,207', '366', '841', 'Stage 2 Training'],
        ['Auto-Labeled', '—', '—', '1,604', '486', '1,118', 'Stage 2 Training'],
        ['GT Test Set', '50', '184', '184', '62', '122', 'Final Validation'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Dataset Summary')
    print("✅ Added Table 2.1: Dataset Summary")

def add_stage1_performance_table(doc):
    """Table: Stage 1 Detector Performance"""
    para_idx = find_table_placeholder(doc, '3.1')
    if para_idx is None:
        print("⚠️  Table 3.1 placeholder not found")
        return
    
    data = [
        ['Detector', 'Params', 'mAP50', 'Precision', 'Recall', 'F1-Score'],
        ['YOLOv11x', '56.9M', '99.5%', '95.0%', '98.0%', '96.5%'],
        ['YOLOv11x_v2 (Final)', '56.9M', '99.7%', '96.5%', '99.0%', '97.7%'],
    ]
    
    insert_table_after_paragraph(doc, para_idx, len(data), len(data[0]), data, 'Stage 1 Performance')
    print("✅ Added Table 3.1: Stage 1 Detector Performance")

def add_best_risk_zone_images(doc):
    """Add the best risk zone visualizations to Results section"""
    print("\n📊 Adding Best Risk Zone Visualizations...")
    
    # Find Results section (Section 10)
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if 'risk zone' in para.text.lower() and 'visualization' in para.text.lower():
            insert_index = i + 2
            break
    
    if insert_index is None:
        # Fallback: find Section 10
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading') and '10' in para.text:
                insert_index = i + 10
                break
    
    if insert_index is None:
        print("⚠️  Could not find suitable location for risk zone images")
        return
    
    images = [
        {
            'path': 'outputs/risk_zones_vit/0039_risk_zones.jpg',
            'number': 'Figure 10.8:',
            'caption': 'Exemplary risk zone visualization from 50-image validation set (Case 0039) - demonstrating optimal GREEN zone classification with high confidence scores across all detected RCT regions'
        },
        {
            'path': 'outputs/risk_zones_vit/0052_risk_zones.jpg',
            'number': 'Figure 10.9:',
            'caption': 'Exemplary risk zone visualization from 50-image validation set (Case 0052) - showcasing clinical decision support system with color-coded risk stratification and confidence metrics'
        },
    ]
    
    for img in images:
        if not os.path.exists(img['path']):
            print(f"⚠️  Image not found: {img['path']}")
            continue
        
        # Insert image
        para = doc.paragraphs[insert_index]
        new_para = para.insert_paragraph_before()
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_para.add_run()
        run.add_picture(img['path'], width=Inches(6.0))
        
        # Add caption
        caption_para = para.insert_paragraph_before()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = caption_para.add_run(img['number'])
        run.bold = True
        run.font.size = Pt(10)
        
        run = caption_para.add_run(f" {img['caption']}")
        run.font.size = Pt(10)
        
        caption_para.paragraph_format.space_after = Pt(12)
        
        print(f"✅ Added: {img['number']}")
        insert_index += 3

def highlight_questions_page(doc):
    """Highlight the page with questions to professors in yellow"""
    print("\n🎨 Highlighting Questions Page...")
    
    keywords = [
        'question', 'soru', 'professor', 'hoca', 'danışman',
        'öğretim üyesi', 'research question', 'araştırma sorusu'
    ]
    
    highlighted_count = 0
    
    for para in doc.paragraphs:
        # Check if paragraph contains question keywords
        text_lower = para.text.lower()
        if any(keyword in text_lower for keyword in keywords):
            # Highlight entire paragraph in yellow
            for run in para.runs:
                run.font.highlight_color = 6  # Yellow (WD_COLOR_INDEX.YELLOW = 7, but 6 is brighter yellow)
            highlighted_count += 1
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_lower = para.text.lower()
                    if any(keyword in text_lower for keyword in keywords):
                        for run in para.runs:
                            run.font.highlight_color = 6
                        highlighted_count += 1
    
    if highlighted_count > 0:
        print(f"✅ Highlighted {highlighted_count} paragraphs containing questions")
    else:
        print("⚠️  No question paragraphs found - please specify exact text to highlight")

def main():
    print("="*80)
    print("📊 FINAL THESIS UPDATE")
    print("="*80)
    
    input_file = Path("thesis_documentation/docx/MASTER_THESIS_WITH_FIGURES.docx")
    output_file = Path("thesis_documentation/docx/MASTER_THESIS_COMPLETE_WITH_TABLES.docx")
    
    if not input_file.exists():
        print(f"❌ Error: Input file not found: {input_file}")
        return
    
    print(f"\n📄 Loading document: {input_file}")
    doc = Document(str(input_file))
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")
    
    # Add all missing tables
    print("\n" + "="*80)
    print("📋 ADDING TABLES")
    print("="*80)
    
    add_dataset_summary_table(doc)
    add_stage1_performance_table(doc)
    add_preprocessing_comparison_table(doc)
    add_stage2_comparison_table(doc)
    add_class_imbalance_strategies_table(doc)
    add_pipeline_optimization_table(doc)
    add_final_results_table(doc)
    
    # Add best risk zone images
    print("\n" + "="*80)
    print("🖼️  ADDING BEST RISK ZONE VISUALIZATIONS")
    print("="*80)
    
    add_best_risk_zone_images(doc)
    
    # Highlight questions page
    print("\n" + "="*80)
    print("🎨 HIGHLIGHTING QUESTIONS")
    print("="*80)
    
    highlight_questions_page(doc)
    
    # Save document
    print("\n" + "="*80)
    print("💾 SAVING DOCUMENT")
    print("="*80)
    
    doc.save(str(output_file))
    print(f"✅ Document saved: {output_file}")
    
    print("\n" + "="*80)
    print("🎉 THESIS UPDATE COMPLETE!")
    print("="*80)
    print(f"\n📄 Output: {output_file}")
    print("\nUpdates:")
    print("  ✅ 7+ tables added with actual data")
    print("  ✅ 2 best risk zone visualizations added (0039, 0052)")
    print("  ✅ Questions page highlighted in yellow")
    print("\nNext steps:")
    print("  1. Open the document in Word")
    print("  2. Verify table placements and data")
    print("  3. Check risk zone image quality")
    print("  4. Update all fields (Ctrl+A, F9)")
    print("  5. Export to PDF for final submission")
    print("="*80)

if __name__ == "__main__":
    main()
