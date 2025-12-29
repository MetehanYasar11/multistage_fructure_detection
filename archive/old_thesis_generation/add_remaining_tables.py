"""
Add remaining tables to specific sections by finding headings.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_table_style(table):
    """Apply professional table styling"""
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row formatting
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4472C4')
        cell._element.get_or_add_tcPr().append(shading_elm)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

def find_heading_by_keywords(doc, keywords):
    """Find heading paragraph by keywords"""
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            text_lower = para.text.lower()
            if any(kw.lower() in text_lower for kw in keywords):
                return i
    return None

def insert_table_after_heading(doc, heading_idx, rows, cols, data):
    """Insert table after a heading"""
    # Create table
    table = doc.add_table(rows=rows, cols=cols)
    
    # Fill data
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_data)
    
    # Style
    add_table_style(table)
    
    # Insert after heading (skip 1-2 paragraphs for intro text)
    insert_para = doc.paragraphs[heading_idx + 3]
    insert_para._element.addprevious(table._element)
    
    return table

def add_missing_tables(doc):
    """Add all missing tables to appropriate sections"""
    
    tables_added = 0
    
    # Table 2.1: Dataset Summary
    print("\n🔍 Looking for Dataset section...")
    idx = find_heading_by_keywords(doc, ['Dataset', 'Data Collection', '2.'])
    if idx:
        data = [
            ['Dataset', 'Images', 'Annotations', 'Crops', 'Fractured', 'Healthy', 'Usage'],
            ['Kaggle', '~3000', '—', '—', '—', '—', 'Stage 1 Training'],
            ['Dataset_2021', '487', '915', '—', '—', '—', 'Stage 1 Validation'],
            ['Manual', '—', '—', '1,207', '366', '841', 'Stage 2 Training'],
            ['Auto-Labeled', '—', '—', '1,604', '486', '1,118', 'Stage 2 Training'],
            ['GT Test', '50', '184', '184', '62', '122', 'Final Validation'],
        ]
        insert_table_after_heading(doc, idx, len(data), len(data[0]), data)
        print("✅ Added Table 2.1: Dataset Summary")
        tables_added += 1
    
    # Table 3.1: Stage 1 Performance
    print("\n🔍 Looking for Stage 1 section...")
    idx = find_heading_by_keywords(doc, ['Stage 1', 'RCT Detection', 'YOLO', '3.'])
    if idx:
        data = [
            ['Detector', 'Parameters', 'mAP50', 'Precision', 'Recall', 'F1-Score'],
            ['YOLOv11x', '56.9M', '99.5%', '95.0%', '98.0%', '96.5%'],
            ['YOLOv11x_v2', '56.9M', '99.7%', '96.5%', '99.0%', '97.7%'],
        ]
        insert_table_after_heading(doc, idx, len(data), len(data[0]), data)
        print("✅ Added Table 3.1: Stage 1 Detector Performance")
        tables_added += 1
    
    # Table 4.1: Preprocessing Comparison
    print("\n🔍 Looking for Preprocessing section...")
    idx = find_heading_by_keywords(doc, ['Preprocessing', 'SR', 'CLAHE', '4.'])
    if idx:
        data = [
            ['Strategy', 'Accuracy', 'Precision', 'Recall', 'F1', 'Δ'],
            ['Baseline', '78.81%', '69.39%', '85.71%', '76.67%', '—'],
            ['SR + CLAHE', '83.44%', '73.58%', '92.86%', '82.11%', '+4.63%'],
            ['CLAHE + Gabor', '~30%', '—', '—', '—', 'FAILED'],
            ['Ensemble', '78.26%', '75.00%', '85.71%', '80.00%', '-0.55%'],
        ]
        insert_table_after_heading(doc, idx, len(data), len(data[0]), data)
        print("✅ Added Table 4.1: Preprocessing Strategy Performance")
        tables_added += 1
    
    # Table 6.1: Stage 2 Model Comparison
    print("\n🔍 Looking for Stage 2 / Model Evolution section...")
    idx = find_heading_by_keywords(doc, ['Stage 2', 'Model Evolution', 'ViT', '6.'])
    if idx:
        data = [
            ['Model', 'Params', 'Test Set', 'Accuracy', 'Recall', 'F1'],
            ['ViT-Tiny', '5.7M', '15 GT crops', '93.33%', '85.71%', '92.31%'],
            ['ViT-Small', '22.0M', '231 auto-labeled', '78.26%', '85.71%', '80.00%'],
            ['ViT-Small', '22.0M', '184 GT crops', '84.78%', '88.71%', '79.71%'],
        ]
        insert_table_after_heading(doc, idx, len(data), len(data[0]), data)
        print("✅ Added Table 6.1: Stage 2 Model Comparison")
        tables_added += 1
    
    # Table 7.1: Class Imbalance Solutions
    print("\n🔍 Looking for Class Imbalance section...")
    idx = find_heading_by_keywords(doc, ['Class Imbalance', 'Imbalance', 'Weighted', '7.'])
    if idx:
        data = [
            ['Strategy', 'Fractured Recall', 'Healthy Recall', 'Accuracy', 'Winner'],
            ['Weighted Loss [0.73, 1.57]', '88.71%', '82.35%', '84.78%', '✓'],
            ['Focal Loss', '85.48%', '79.41%', '81.52%', '—'],
            ['SMOTE', '83.87%', '82.35%', '82.61%', '—'],
            ['Balanced Sampling', '87.10%', '76.47%', '80.43%', '—'],
        ]
        insert_table_after_heading(doc, idx, len(data), len(data[0]), data)
        print("✅ Added Table 7.1: Class Imbalance Solutions")
        tables_added += 1
    
    return tables_added

def main():
    print("="*80)
    print("📋 ADDING REMAINING TABLES")
    print("="*80)
    
    input_file = "thesis_documentation/docx/MASTER_THESIS_COMPLETE_WITH_TABLES.docx"
    output_file = "thesis_documentation/docx/MASTER_THESIS_FINAL_COMPLETE.docx"
    
    doc = Document(input_file)
    print(f"\n📄 Loaded: {input_file}")
    print(f"   Paragraphs: {len(doc.paragraphs)}")
    print(f"   Existing tables: {len(doc.tables)}")
    
    tables_added = add_missing_tables(doc)
    
    print("\n" + "="*80)
    print("💾 SAVING")
    print("="*80)
    doc.save(output_file)
    print(f"✅ Saved: {output_file}")
    print(f"   New tables added: {tables_added}")
    print(f"   Total tables now: {len(doc.tables)}")
    
    print("\n" + "="*80)
    print("🎉 COMPLETE!")
    print("="*80)

if __name__ == "__main__":
    main()
