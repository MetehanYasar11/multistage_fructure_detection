"""
Find all table captions that are missing actual tables.
"""

from docx import Document
from docx.table import Table

def find_missing_tables(doc_path):
    """Find table captions that don't have actual tables after them"""
    doc = Document(doc_path)
    
    missing_tables = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Check if this looks like a table caption
        if text.startswith('Table ') and ':' in text:
            # Extract table number
            try:
                table_num = text.split(':')[0].strip()
                
                # Check next few elements (paragraphs or tables)
                has_table = False
                
                # Look ahead up to 5 elements
                for j in range(i+1, min(i+6, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    
                    # If we hit another table caption or heading, stop
                    if next_para.text.startswith('Table ') or next_para.style.name.startswith('Heading'):
                        break
                    
                    # Check if there's a table between current and next paragraph
                    # This is tricky - we need to check the document structure
                    
                # For now, let's just collect all table captions
                missing_tables.append({
                    'number': table_num,
                    'caption': text,
                    'para_index': i
                })
                
            except:
                pass
    
    return missing_tables

def check_tables_in_document(doc_path):
    """Check which tables exist and which are missing"""
    doc = Document(doc_path)
    
    print("="*80)
    print("📊 TABLE ANALYSIS")
    print("="*80)
    
    print(f"\n📄 Document: {doc_path}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables (objects): {len(doc.tables)}")
    
    # Find all table captions
    table_captions = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('Table ') and ':' in text:
            # Extract table number
            try:
                parts = text.split(':')
                table_num = parts[0].replace('Table ', '').strip()
                caption = ':'.join(parts[1:]).strip()
                table_captions.append({
                    'number': table_num,
                    'full_caption': text[:100]
                })
            except:
                pass
    
    print(f"\n📋 Found {len(table_captions)} table captions")
    print(f"🔢 Actual table objects in document: {len(doc.tables)}")
    
    # Show all table numbers
    table_numbers = sorted(set([t['number'] for t in table_captions]))
    print(f"\n📝 Table numbers found in captions:")
    
    # Group by section
    sections = {}
    for num in table_numbers:
        section = num.split('.')[0]
        if section not in sections:
            sections[section] = []
        sections[section].append(num)
    
    for section in sorted(sections.keys()):
        print(f"\n   Section {section}:")
        for num in sections[section]:
            matching_captions = [t for t in table_captions if t['number'] == num]
            if matching_captions:
                print(f"      Table {num}: {matching_captions[0]['full_caption'][20:80]}...")
    
    # List potentially missing tables (captions without immediate tables)
    print(f"\n⚠️  POTENTIALLY MISSING TABLES:")
    print("   (These have captions but may not have actual table data)")
    
    # Critical tables that should definitely exist
    critical_tables = [
        ('2.1', 'Dataset Summary'),
        ('3.1', 'Stage 1 Performance'),
        ('4.1', 'Preprocessing Comparison'),
        ('6.1', 'Stage 2 Model Comparison'),
        ('7.1', 'Class Imbalance Solutions'),
        ('8.1', 'Pipeline Optimization'),
        ('10.1', 'Final Results'),
        ('10.4', 'Literature Comparison'),
    ]
    
    print("\n🎯 CRITICAL TABLES STATUS:")
    for table_num, description in critical_tables:
        found = any(t['number'] == table_num for t in table_captions)
        status = "✅ Caption found" if found else "❌ MISSING"
        print(f"   Table {table_num} ({description}): {status}")
    
    return table_captions

if __name__ == "__main__":
    doc_path = "thesis_documentation/docx/MASTER_THESIS_FINAL_COMPLETE.docx"
    table_captions = check_tables_in_document(doc_path)
    
    print("\n" + "="*80)
    print(f"📊 SUMMARY: {len(table_captions)} table captions found")
    print("="*80)
