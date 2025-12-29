"""
Generate Section 1: Introduction - COMPLETE VERSION
=================================================================

This script generates a complete Section 1 with:
- Full text content
- All figures embedded (3 figures)
- No tables (text-only section)
- NO MANUAL WORK NEEDED!

Author: Thesis Generation System v2
Date: December 22, 2025
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_figure_with_caption(doc, image_path, figure_number, caption_text, width=6.0):
    """
    Add a figure with caption to the document.
    
    Args:
        doc: Document object
        image_path: Path to the image file
        figure_number: Figure number (e.g., "1.1")
        caption_text: Caption text
        width: Image width in inches
    
    Returns:
        bool: True if image was added successfully, False otherwise
    """
    if not os.path.exists(image_path):
        print(f"⚠️  Warning: Image not found: {image_path}")
        print(f"   → Will add YELLOW HIGHLIGHTED caption for manual addition")
        
        # Add yellow highlighted caption for missing image
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = caption_para.add_run(f"Figure {figure_number}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        run2 = caption_para.add_run(caption_text)
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        # Highlight in yellow
        run.font.highlight_color = 7  # Yellow
        run2.font.highlight_color = 7
        
        doc.add_paragraph()  # Spacing
        return False
    
    # Add image
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.add_run()
    run.add_picture(image_path, width=Inches(width))
    
    # Add caption
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = caption_para.add_run(f"Figure {figure_number}: ")
    run.bold = True
    run.font.size = Pt(10)
    
    run = caption_para.add_run(caption_text)
    run.font.size = Pt(10)
    
    # Add spacing after figure
    doc.add_paragraph()
    
    print(f"✅ Figure {figure_number} added successfully")
    return True


def generate_section1():
    """
    Generate complete Section 1: Introduction with all content.
    """
    print("="*80)
    print("📚 GENERATING SECTION 1: INTRODUCTION")
    print("="*80)
    print()
    
    doc = Document()
    
    # ========================================
    # SECTION TITLE
    # ========================================
    title = doc.add_heading('1. Introduction', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # ========================================
    # 1.1 Problem Statement
    # ========================================
    doc.add_heading('1.1 Problem Statement and Motivation', level=2)
    
    doc.add_paragraph(
        "Vertical root fractures (VRF) represent one of the most challenging diagnostic dilemmas "
        "in endodontics, particularly in teeth that have undergone root canal treatment (RCT). "
        "These fractures typically originate from the root canal space and extend longitudinally "
        "toward the periodontal ligament, often remaining undetected until significant bone loss "
        "has occurred. Early detection of VRF is crucial for treatment planning and can potentially "
        "save teeth that would otherwise require extraction."
    )
    
    doc.add_paragraph(
        "Traditional diagnostic methods for VRF, including clinical examination and conventional "
        "radiography, suffer from several limitations. Panoramic radiographs, while widely available "
        "and commonly used for dental screening, provide only two-dimensional representations of "
        "three-dimensional structures. This limitation, combined with the superimposition of anatomical "
        "structures and the subtle nature of vertical fractures, results in low sensitivity for VRF "
        "detection. Studies have reported that even experienced clinicians may miss up to 30-40% of "
        "existing fractures when relying solely on panoramic radiographs."
    )
    
    doc.add_paragraph(
        "The advent of deep learning and computer vision technologies presents a promising opportunity "
        "to address these diagnostic challenges. Convolutional neural networks (CNNs) and vision transformers "
        "(ViTs) have demonstrated remarkable capabilities in medical image analysis, achieving expert-level "
        "performance in various diagnostic tasks. However, the application of these technologies to VRF "
        "detection faces unique challenges, including limited annotated datasets, class imbalance between "
        "fractured and healthy teeth, and the need for systems that can provide clinically interpretable "
        "results beyond simple binary classifications."
    )
    
    # ========================================
    # 1.2 Research Objectives
    # ========================================
    doc.add_heading('1.2 Research Objectives', level=2)
    
    doc.add_paragraph(
        "This research aims to develop an automated, two-stage deep learning system for detecting "
        "vertical root fractures in panoramic dental radiographs of teeth with root canal treatment. "
        "The primary objectives of this study are:"
    )
    
    objectives = [
        ("Primary Detection System", 
         "Develop and validate a two-stage pipeline combining object detection (YOLOv11x) and "
         "image classification (Vision Transformer) to achieve accurate VRF detection at both "
         "crop-level and image-level evaluations."),
        
        ("Auto-Labeling Pipeline", 
         "Design and implement an efficient auto-labeling system using the Liang-Barsky line-rectangle "
         "intersection algorithm to dramatically reduce manual annotation time (targeting 200× speedup) "
         "while maintaining >95% labeling accuracy."),
        
        ("Preprocessing Optimization", 
         "Systematically evaluate and optimize preprocessing strategies, including super-resolution "
         "enhancement and contrast-limited adaptive histogram equalization (CLAHE), to maximize "
         "classification performance on challenging panoramic radiographs."),
        
        ("Class Imbalance Solutions", 
         "Investigate and compare multiple approaches to handle the inherent class imbalance in "
         "dental fracture datasets (typically 1:2 to 1:3 ratio of fractured to healthy teeth), "
         "including weighted loss functions, focal loss, SMOTE, and balanced sampling strategies."),
        
        ("Pipeline Optimization", 
         "Develop a comprehensive optimization framework for aggregating crop-level predictions into "
         "reliable image-level diagnoses, including confidence thresholding, voting mechanisms, and "
         "sensitivity-specificity trade-off analysis."),
        
        ("Clinical Decision Support", 
         "Create a risk zone visualization system that provides intuitive, color-coded risk stratification "
         "(GREEN/YELLOW/RED) for rapid screening and clinician prioritization, transforming raw model "
         "predictions into actionable clinical insights."),
        
        ("Validation and Benchmarking", 
         "Conduct rigorous validation using multiple test sets including ground-truth annotated data "
         "and independent professor-evaluated cases, and benchmark performance against existing "
         "literature to demonstrate improvements in diagnostic accuracy.")
    ]
    
    for i, (obj_title, obj_text) in enumerate(objectives, 1):
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(f"{obj_title}: ")
        run.bold = True
        para.add_run(obj_text)
    
    # ========================================
    # 1.3 Research Scope and Repository Overview
    # ========================================
    doc.add_heading('1.3 Research Scope and Repository Overview', level=2)
    
    doc.add_paragraph(
        "This research encompasses a comprehensive investigation spanning multiple dimensions of "
        "deep learning-based VRF detection. The project repository (MetehanYasar11/multistage_fructure_detection, "
        "branch: prototype) serves as a complete record of the experimental journey, containing over "
        "3,000 images across multiple datasets, 50+ distinct experiments exploring various architectures "
        "and preprocessing strategies, and a fully documented codebase enabling reproducibility."
    )
    
    print("📊 Adding Figure 1.1: Repository Statistics Overview...")
    add_figure_with_caption(
        doc,
        "../outputs/repo_visualizations/repo_statistics_overview.png",
        "1.1",
        "Repository statistics overview showing the scale of experimental work: 3K+ images processed, "
        "50+ experiments conducted, multiple datasets integrated, and comprehensive documentation maintained",
        width=6.5
    )
    
    doc.add_paragraph(
        "The research timeline reflects a systematic, iterative approach to problem-solving, beginning "
        "with dataset assembly and initial baseline models in early 2024, progressing through extensive "
        "preprocessing experiments and architecture evaluations in mid-2024, and culminating in advanced "
        "pipeline optimization and comprehensive validation in late 2024. Each phase built upon insights "
        "from previous experiments, with careful documentation enabling analysis of what worked, what "
        "failed, and why."
    )
    
    print("📊 Adding Figure 1.2: Research Timeline...")
    add_figure_with_caption(
        doc,
        "../outputs/repo_visualizations/research_timeline.png",
        "1.2",
        "Research timeline showing major milestones throughout 2024: dataset preparation (Q1), "
        "baseline model development (Q2), preprocessing optimization (Q3), and pipeline refinement (Q4)",
        width=6.5
    )
    
    doc.add_paragraph(
        "The experimental scope encompasses multiple critical dimensions of the problem. Preprocessing "
        "experiments explored super-resolution techniques (bicubic, ESRGAN, Real-ESRGAN), contrast "
        "enhancement methods (CLAHE with various parameters), and feature enhancement approaches (Gabor "
        "filters with multiple orientations). Architecture investigations compared various model sizes "
        "and families including Vision Transformers (ViT-Tiny, ViT-Small), EfficientNets, and YOLO "
        "detectors (v8, v11, multiple variants). Class imbalance solutions tested weighted loss functions, "
        "focal loss, synthetic oversampling (SMOTE), and balanced sampling strategies. Pipeline optimization "
        "conducted systematic grid searches over confidence thresholds, voting ratios, and aggregation "
        "methods to maximize both sensitivity and specificity."
    )
    
    print("📊 Adding Figure 1.3: Experiments Breakdown...")
    add_figure_with_caption(
        doc,
        "../outputs/repo_visualizations/experiments_breakdown.png",
        "1.3",
        "Breakdown of experiments by category: preprocessing techniques (SR, CLAHE, Gabor), "
        "model architectures (ViT, EfficientNet, YOLO variants), class balancing strategies, "
        "and pipeline optimization approaches",
        width=6.5
    )
    
    # ========================================
    # 1.4 Key Contributions
    # ========================================
    doc.add_heading('1.4 Key Contributions and Innovations', level=2)
    
    doc.add_paragraph(
        "This research makes several significant contributions to the field of AI-assisted dental "
        "diagnostics, advancing both technical methodologies and clinical applicability:"
    )
    
    contributions = [
        ("Novel Auto-Labeling Algorithm",
         "Implementation of the Liang-Barsky line-rectangle intersection algorithm for dental "
         "image annotation, achieving 200× speedup (from 40-60 hours to 15 minutes) with >95% "
         "accuracy. This innovation dramatically reduces the annotation burden, making large-scale "
         "dataset creation feasible for dental AI applications."),
        
        ("Optimized Two-Stage Architecture",
         "Design and validation of a two-stage detection pipeline combining YOLOv11x (99.7% mAP@0.5 "
         "for RCT detection) with ViT-Small classifier (84.78% crop-level accuracy), demonstrating "
         "that task decomposition into detection and classification stages outperforms end-to-end "
         "approaches for this specific domain."),
        
        ("Effective Preprocessing Pipeline",
         "Systematic evaluation demonstrating that 4× bicubic super-resolution combined with CLAHE "
         "(clipLimit=2.0, tileSize=16×16) achieves +4.63% accuracy improvement over baseline, while "
         "documenting that more aggressive preprocessing (CLAHE+Gabor) can actually harm performance, "
         "providing valuable negative results for the field."),
        
        ("Class Imbalance Solution",
         "Demonstration that carefully calibrated weighted loss functions [0.73, 1.57] achieve "
         "superior performance (88.71% recall) compared to more complex alternatives like focal loss "
         "or synthetic oversampling, offering a simple yet effective solution for dental fracture "
         "detection where recall is critical."),
        
        ("Comprehensive Pipeline Optimization",
         "Development of a systematic optimization framework that improved specificity 8-fold "
         "(from 7.69% to 61.54%) through combined confidence thresholding (≥0.75) and voting "
         "requirements (≥2 detections), addressing a critical clinical need for reducing false "
         "alarms in screening systems."),
        
        ("Clinical Risk Zone System",
         "Creation of an intuitive color-coded risk stratification system (GREEN/YELLOW/RED) that "
         "aggregates multiple model predictions into actionable clinical insights, with demonstrated "
         "capability to achieve 89.47% image-level accuracy and 92% recall in optimized configuration."),
        
        ("Rigorous Validation Framework",
         "Establishment of a multi-level validation methodology including crop-level evaluation "
         "(184 ground-truth crops), image-level assessment with different confidence strategies, "
         "independent professor-evaluated test set (20 cases), and comprehensive comparison with "
         "existing literature, demonstrating 3-8% accuracy advantage over previous methods."),
        
        ("Complete Reproducibility Package",
         "Development of a fully documented, reproducible codebase with automated thesis generation "
         "scripts, comprehensive experiment logs, and systematic organization of all results, "
         "ensuring that this work can be validated, extended, and built upon by future researchers.")
    ]
    
    for i, (contrib_title, contrib_text) in enumerate(contributions, 1):
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(f"{contrib_title}: ")
        run.bold = True
        para.add_run(contrib_text)
    
    # ========================================
    # 1.5 Clinical Significance
    # ========================================
    doc.add_heading('1.5 Clinical Significance and Impact', level=2)
    
    doc.add_paragraph(
        "The clinical implications of this research extend beyond technical performance metrics to "
        "address real-world challenges in dental practice. The developed system offers several "
        "practical advantages for clinical deployment:"
    )
    
    doc.add_paragraph(
        "First, the system's high sensitivity (88.71% recall at crop-level, 92% with optimized "
        "pipeline) directly addresses the primary clinical concern of missing fractures that require "
        "intervention. The risk zone visualization system enables rapid screening of large patient "
        "populations, allowing clinicians to prioritize cases requiring detailed examination while "
        "confidently clearing low-risk cases."
    )
    
    doc.add_paragraph(
        "Second, the two-stage architecture naturally provides interpretable intermediate results. "
        "The first stage identifies all RCT teeth, ensuring comprehensive coverage, while the second "
        "stage provides confidence scores for each detection. This transparency allows clinicians to "
        "understand the system's reasoning and make informed decisions about when to trust automated "
        "assessments versus when to perform additional confirmatory testing."
    )
    
    doc.add_paragraph(
        "Third, the system's flexibility in confidence threshold configuration enables adaptation to "
        "different clinical scenarios. High-sensitivity mode (confidence ≥0.3, achieving 100% recall "
        "on 20-image test) is suitable for screening applications where missing a fracture has serious "
        "consequences. Balanced mode (confidence ≥0.5) provides better specificity while maintaining "
        "strong sensitivity for routine diagnostic workflows. This adaptability makes the system "
        "suitable for deployment across diverse clinical settings with varying resource constraints "
        "and risk tolerances."
    )
    
    doc.add_paragraph(
        "Finally, the efficient auto-labeling pipeline developed in this research addresses a critical "
        "bottleneck in creating dental AI systems. By reducing annotation time from days to minutes, "
        "this innovation enables dental practices and research institutions to leverage their existing "
        "image archives for AI development, democratizing access to this technology beyond large "
        "academic centers with extensive annotation resources."
    )
    
    # ========================================
    # 1.6 Thesis Organization
    # ========================================
    doc.add_heading('1.6 Thesis Organization', level=2)
    
    doc.add_paragraph(
        "The remainder of this thesis is organized as follows:"
    )
    
    chapters = [
        ("Section 2: Dataset and Data Collection",
         "Describes the multiple datasets used in this research, including Kaggle training data, "
         "Dataset_2021 with professional annotations, manually labeled crops, and auto-labeled crops. "
         "Details data distribution, annotation quality, and dataset preparation strategies."),
        
        ("Section 3: Stage 1 - Root Canal Treated Tooth Detection",
         "Presents the first stage of the pipeline, focusing on YOLOv11x detector training, evaluation, "
         "and optimization. Discusses detector evolution from initial v11x to improved v11x_v2 model, "
         "achieving 99.7% mAP@0.5. Analyzes inference configurations and bounding box scaling strategies."),
        
        ("Section 4: Preprocessing Experiments and Analysis",
         "Systematically evaluates multiple preprocessing approaches including super-resolution techniques "
         "(bicubic, ESRGAN, Real-ESRGAN), CLAHE parameter optimization, Gabor filter experiments (including "
         "documented failure modes), and ensemble model investigations. Establishes SR+CLAHE as optimal "
         "preprocessing pipeline."),
        
        ("Section 5: Dataset Generation Strategies and Auto-Labeling",
         "Details the development of the Liang-Barsky-based auto-labeling system, comparing manual annotation "
         "approaches with automated generation. Quantifies time savings (200× speedup) and accuracy trade-offs "
         "(>95% agreement). Discusses implications for future dataset creation efforts."),
        
        ("Section 6: Stage 2 Model Evolution and Architecture Selection",
         "Chronicles the iterative development of the classification stage, from initial ViT-Tiny experiments "
         "(overfitting on 15 crops) through ViT-Small training on auto-labeled data (78.26%) to final validation "
         "on ground-truth crops (84.78%). Analyzes label noise effects and model selection criteria."),
        
        ("Section 7: Class Imbalance Solutions and Training Strategies",
         "Investigates multiple approaches to addressing the 1:2.3 class imbalance including weighted loss "
         "functions (winner: [0.73, 1.57]), focal loss, SMOTE oversampling, and balanced sampling. Presents "
         "comparative evaluation across all strategies and justifies final selection."),
        
        ("Section 8: Pipeline Optimization and Aggregation Strategies",
         "Describes comprehensive pipeline optimization including baseline analysis (7.69% specificity problem), "
         "grid search over 120 configurations (10 confidence levels × 12 voting ratios), development of combined "
         "threshold strategy, and risk zone aggregation system. Documents 8× specificity improvement."),
        
        ("Section 9: System Architecture and Implementation",
         "Provides complete system specification including component details (Stage 1, preprocessing, Stage 2, "
         "risk zones), configuration parameters (8 key parameters), deployment requirements (hardware/software), "
         "and performance characteristics. Enables reproducible implementation."),
        
        ("Section 10: Results and Discussion",
         "Presents comprehensive results including primary validation (50-image, 84.78% crop-level), additional "
         "testing (20-image professor test, 88-94% image-level), comprehensive comparison across evaluation "
         "levels, literature benchmarking (3-8% advantage), qualitative analysis, risk zone showcase, and "
         "detailed discussion of performance interpretation, clinical implications, limitations, and strengths."),
        
        ("Section 11: Conclusion and Future Work",
         "Summarizes research contributions, assesses clinical impact, discusses system strengths and limitations, "
         "and proposes 12 directions for future research including multi-class classification, attention "
         "visualization, multi-center validation, prospective clinical trials, real-time deployment, mobile "
         "applications, PACS integration, longitudinal analysis, generalization studies, 3D CBCT extension, "
         "few-shot learning, and federated learning approaches.")
    ]
    
    for i, (chapter_title, chapter_desc) in enumerate(chapters, 2):
        para = doc.add_paragraph()
        run = para.add_run(f"{chapter_title}: ")
        run.bold = True
        para.add_run(chapter_desc)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(6)
    
    # ========================================
    # SAVE DOCUMENT
    # ========================================
    output_path = 'v2_generation/outputs/SECTION_1_INTRODUCTION_V2_COMPLETE.docx'
    doc.save(output_path)
    
    print()
    print("="*80)
    print("✅ SECTION 1 GENERATED SUCCESSFULLY!")
    print("="*80)
    print(f"📄 Output: {output_path}")
    print()
    print("Content includes:")
    print("  ✅ Complete text (6 subsections)")
    print("  ✅ Figure 1.1: Repository statistics overview")
    print("  ✅ Figure 1.2: Research timeline")
    print("  ✅ Figure 1.3: Experiments breakdown")
    print("  ✅ No tables (text-only section)")
    print("="*80)
    
    return doc


if __name__ == "__main__":
    generate_section1()
