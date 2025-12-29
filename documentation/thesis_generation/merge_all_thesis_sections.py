"""
Merge ALL thesis sections (1-11) into a single master document

This script:
1. Loads all 11 DOCX section files
2. Combines them into a single document
3. Adds proper page breaks between sections
4. Generates Table of Contents placeholder
5. Saves as MASTER_THESIS_COMPLETE.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
import os

def add_cover_page(doc):
    """Add a professional cover page"""
    # University name
    para = doc.add_paragraph()
    run = para.add_run("ISTANBUL TECHNICAL UNIVERSITY")
    run.bold = True
    run.font.size = Pt(16)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # Faculty
    para = doc.add_paragraph()
    run = para.add_run("GRADUATE SCHOOL OF SCIENCE ENGINEERING AND TECHNOLOGY")
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Thesis title
    para = doc.add_paragraph()
    run = para.add_run("AUTOMATED DETECTION AND CLASSIFICATION OF\nROOT CANAL TREATMENT FRACTURES IN\nPANORAMIC DENTAL X-RAYS USING\nVISION TRANSFORMERS")
    run.bold = True
    run.font.size = Pt(18)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Thesis type
    para = doc.add_paragraph()
    run = para.add_run("MASTER'S THESIS")
    run.bold = True
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Author
    para = doc.add_paragraph()
    run = para.add_run("Metehan YAŞAR")
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    para = doc.add_paragraph()
    run = para.add_run("DECEMBER 2024")
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Page break after cover
    doc.add_page_break()

def add_toc_placeholder(doc):
    """Add Table of Contents placeholder"""
    # TOC Title
    para = doc.add_paragraph()
    run = para.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(16)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # TOC note
    para = doc.add_paragraph()
    run = para.add_run("[Table of Contents will be auto-generated in Microsoft Word]")
    run.italic = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("Instructions:")
    run.bold = True
    
    instructions = [
        "1. In Microsoft Word, go to 'References' tab",
        "2. Click 'Table of Contents'",
        "3. Choose a style (Automatic Table 1 or 2)",
        "4. The TOC will automatically populate with all section headings",
        "5. To update: Right-click TOC → Update Field → Update entire table"
    ]
    
    for instruction in instructions:
        para = doc.add_paragraph(instruction)
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()

def add_abstract(doc):
    """Add abstract page"""
    # Abstract title
    para = doc.add_paragraph()
    run = para.add_run("ABSTRACT")
    run.bold = True
    run.font.size = Pt(16)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # Abstract content
    abstract_text = (
        "Vertical root fractures (VRF) in root canal treated teeth are critical pathologies that "
        "significantly impact treatment outcomes and tooth survival. Traditional manual inspection "
        "of panoramic dental X-rays is time-consuming, subjective, and prone to inter-observer variability. "
        "This thesis presents an automated deep learning pipeline for detection and classification of "
        "root canal treatment (RCT) fractures using Vision Transformer architecture.\n\n"
        
        "The proposed system employs a two-stage approach: (1) Stage 1 uses YOLOv11x object detector "
        "for RCT region extraction from full panoramic X-rays, achieving 95% precision and 98% recall; "
        "(2) Stage 2 applies Vision Transformer Small (ViT-Small) for binary classification (Healthy vs "
        "Fractured) of individual RCT crops. Novel contributions include Super-Resolution combined with "
        "CLAHE preprocessing (+4.63% accuracy improvement), weighted cross-entropy loss for severe class "
        "imbalance (38.89% → 88.71% recall improvement), and an automated labeling system using Liang-Barsky "
        "line-clipping algorithm (200× speedup with >95% accuracy).\n\n"
        
        "The system was validated on multiple datasets: (1) 50-image primary validation achieving "
        "84.78% crop-level accuracy (88.71% recall, 72.37% precision) on 184 RCT crops, and (2) 20-image "
        "clinical test achieving 88.24-94.44% image-level accuracy. A comprehensive analysis revealed "
        "Stage 1 detector sensitivity to image source distribution shift, with five contributing factors "
        "identified and mitigation strategies proposed. The risk zone visualization system (GREEN/YELLOW/RED) "
        "provides intuitive clinical decision support, estimated to reduce radiologist review time by 30-40%.\n\n"
        
        "Results demonstrate that Vision Transformers outperform CNN baselines (ViT: 87.96% vs "
        "EfficientNet: 85.74% vs ResNet: 83.72%) for dental fracture detection. The system is positioned "
        "as a clinical decision support tool rather than autonomous diagnosis, with high recall prioritizing "
        "patient safety by minimizing missed fractures. Data integrity was rigorously validated with no "
        "data leakage confirmed. The system is ready for prospective clinical validation at multiple "
        "institutions, with deployment recommendations including confidence threshold tuning (conf=0.5) "
        "and optional fine-tuning on 50-100 local images.\n\n"
        
        "This research establishes Vision Transformers as a viable architecture for dental radiography "
        "AI, with methodological insights generalizable to other medical imaging tasks involving class "
        "imbalance, limited training data, and subtle visual patterns."
    )
    
    para = doc.add_paragraph(abstract_text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_paragraph()
    
    # Keywords
    para = doc.add_paragraph()
    run = para.add_run("Keywords: ")
    run.bold = True
    run = para.add_run(
        "Dental X-ray Analysis, Vertical Root Fracture Detection, Vision Transformer, "
        "Deep Learning, Medical Image Classification, Class Imbalance, Super-Resolution, "
        "CLAHE Preprocessing, Weighted Loss Function, Clinical Decision Support"
    )
    
    doc.add_page_break()

def merge_sections():
    """Merge all thesis sections into a single document"""
    
    print("=" * 80)
    print("📚 MERGING ALL THESIS SECTIONS (1-11)")
    print("=" * 80)
    print()
    
    # Create new master document
    master_doc = Document()
    
    # Add cover page
    print("✅ Adding cover page...")
    add_cover_page(master_doc)
    
    # Add TOC placeholder
    print("✅ Adding Table of Contents placeholder...")
    add_toc_placeholder(master_doc)
    
    # Add abstract
    print("✅ Adding abstract...")
    add_abstract(master_doc)
    
    # List of section files in order
    section_files = [
        'THESIS_SECTION_1_INTRODUCTION.docx',
        'THESIS_SECTION_2_LITERATURE_REVIEW.docx',
        'THESIS_SECTION_3_DATASET_PREPROCESSING.docx',
        'THESIS_SECTION_4_METHODOLOGY.docx',
        'THESIS_SECTION_5_IMPLEMENTATION.docx',
        'THESIS_SECTION_6_EXPERIMENTS.docx',
        'THESIS_SECTION_7_AUTOLABELING.docx',
        'THESIS_SECTION_8_PIPELINE_OPTIMIZATION.docx',
        'THESIS_SECTION_9_FINAL_ARCHITECTURE.docx',
        'THESIS_SECTION_10_RESULTS_DISCUSSION.docx',
        'THESIS_SECTION_11_CONCLUSION_FUTURE_WORK.docx'
    ]
    
    print()
    print("📄 Merging sections...")
    print("-" * 80)
    
    total_paragraphs = 0
    
    for i, section_file in enumerate(section_files, 1):
        if not os.path.exists(section_file):
            print(f"⚠️  WARNING: {section_file} not found, skipping...")
            continue
        
        print(f"   Section {i:2d}: {section_file:50s} ", end="")
        
        try:
            # Load section document
            section_doc = Document(section_file)
            
            # Count paragraphs
            para_count = len(section_doc.paragraphs)
            total_paragraphs += para_count
            
            # Copy all paragraphs from section to master
            for para in section_doc.paragraphs:
                # Create new paragraph in master doc
                new_para = master_doc.add_paragraph()
                
                # Copy paragraph style
                new_para.style = para.style
                new_para.alignment = para.alignment
                
                # Copy all runs (to preserve formatting)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
            
            # Copy tables
            for table in section_doc.tables:
                # Create new table in master
                new_table = master_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style
                
                # Copy table content
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
                        # Copy cell formatting
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.bold or run.italic:
                                    for new_para in new_table.rows[i].cells[j].paragraphs:
                                        for new_run in new_para.runs:
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
            
            # Add page break after each section (except last)
            if i < len(section_files):
                master_doc.add_page_break()
            
            print(f"✅ ({para_count:4d} paragraphs)")
            
        except Exception as e:
            print(f"❌ ERROR: {str(e)}")
            continue
    
    print("-" * 80)
    print(f"✅ Total paragraphs merged: {total_paragraphs}")
    print()
    
    # Save master document
    output_file = 'MASTER_THESIS_COMPLETE.docx'
    master_doc.save(output_file)
    
    # Get file size
    file_size = os.path.getsize(output_file) / (1024 * 1024)  # MB
    
    print("=" * 80)
    print("✅ MASTER THESIS DOCUMENT CREATED SUCCESSFULLY!")
    print("=" * 80)
    print(f"Output file: {output_file}")
    print(f"File size: {file_size:.2f} MB")
    print(f"Total sections merged: {len(section_files)}")
    print(f"Total paragraphs: {total_paragraphs}")
    print()
    print("📋 Document Structure:")
    print("   • Cover Page")
    print("   • Table of Contents (placeholder - auto-generate in Word)")
    print("   • Abstract")
    print("   • Section 1: Introduction")
    print("   • Section 2: Literature Review")
    print("   • Section 3: Dataset & Preprocessing")
    print("   • Section 4: Methodology")
    print("   • Section 5: Implementation")
    print("   • Section 6: Experiments")
    print("   • Section 7: Auto-Labeling")
    print("   • Section 8: Pipeline Optimization")
    print("   • Section 9: System Architecture")
    print("   • Section 10: Results & Discussion")
    print("   • Section 11: Conclusion & Future Work")
    print()
    print("📌 Next Steps:")
    print("   1. Open MASTER_THESIS_COMPLETE.docx in Microsoft Word")
    print("   2. Go to References → Table of Contents → Auto Table")
    print("   3. Add page numbers (Insert → Page Number)")
    print("   4. Search outputs/ and runs/ for PNG visualizations")
    print("   5. Embed figures with captions")
    print("   6. Final proofread and formatting")
    print("   7. Export to PDF for submission")
    print("=" * 80)

if __name__ == "__main__":
    merge_sections()
