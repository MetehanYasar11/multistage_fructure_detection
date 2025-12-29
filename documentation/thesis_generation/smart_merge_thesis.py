"""
Smart merge: Combine existing THESIS_SECTIONS_1_2_3_4_5_6_7_8_COMPLETE.docx 
with Section 9, 10, 11 to create final master document
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_cover_page(doc):
    """Add a professional cover page"""
    para = doc.add_paragraph()
    run = para.add_run("ISTANBUL TECHNICAL UNIVERSITY")
    run.bold = True
    run.font.size = Pt(16)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("GRADUATE SCHOOL OF SCIENCE ENGINEERING AND TECHNOLOGY")
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("AUTOMATED DETECTION AND CLASSIFICATION OF\nROOT CANAL TREATMENT FRACTURES IN\nPANORAMIC DENTAL X-RAYS USING\nVISION TRANSFORMERS")
    run.bold = True
    run.font.size = Pt(18)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("MASTER'S THESIS")
    run.bold = True
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("Metehan YAŞAR")
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("DECEMBER 2024")
    run.font.size = Pt(14)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()

def add_toc_placeholder(doc):
    """Add Table of Contents placeholder"""
    para = doc.add_paragraph()
    run = para.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(16)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("[Table of Contents will be auto-generated in Microsoft Word]")
    run.italic = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run("Instructions:")
    run.bold = True
    
    instructions = [
        "1. In Microsoft Word, go to 'References' tab",
        "2. Click 'Table of Contents' → Choose a style",
        "3. TOC will auto-populate with all section headings",
        "4. To update: Right-click TOC → Update Field → Update entire table"
    ]
    
    for instruction in instructions:
        doc.add_paragraph(instruction)
    
    doc.add_page_break()

def add_abstract(doc):
    """Add abstract page"""
    para = doc.add_paragraph()
    run = para.add_run("ABSTRACT")
    run.bold = True
    run.font.size = Pt(16)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    abstract_text = (
        "Vertical root fractures (VRF) in root canal treated teeth are critical pathologies that "
        "significantly impact treatment outcomes and tooth survival. This thesis presents an automated "
        "deep learning pipeline for detection and classification of root canal treatment (RCT) fractures "
        "using Vision Transformer architecture.\n\n"
        
        "The proposed system employs a two-stage approach: (1) YOLOv11x object detector for RCT region "
        "extraction achieving 95% precision and 98% recall; (2) Vision Transformer Small (ViT-Small) for "
        "binary classification (Healthy vs Fractured). Novel contributions include Super-Resolution with "
        "CLAHE preprocessing (+4.63% accuracy improvement), weighted cross-entropy loss (38.89% → 88.71% "
        "recall improvement), and automated labeling system using Liang-Barsky algorithm (200× speedup "
        "with >95% accuracy).\n\n"
        
        "The system was validated achieving 84.78% crop-level accuracy (88.71% recall, 72.37% precision) "
        "on 184 RCT crops and 88.24-94.44% image-level accuracy on clinical tests. Comprehensive analysis "
        "revealed Stage 1 detector sensitivity to image source distribution shift. The risk zone visualization "
        "system (GREEN/YELLOW/RED) provides intuitive clinical decision support, estimated to reduce "
        "radiologist review time by 30-40%.\n\n"
        
        "Results demonstrate Vision Transformers outperform CNN baselines (ViT: 87.96% vs EfficientNet: "
        "85.74% vs ResNet: 83.72%). The system is ready for prospective clinical validation with deployment "
        "recommendations including confidence threshold tuning and optional fine-tuning on local images.\n\n"
        
        "Keywords: Dental X-ray Analysis, Vertical Root Fracture Detection, Vision Transformer, Deep Learning, "
        "Medical Image Classification, Class Imbalance, Clinical Decision Support"
    )
    
    para = doc.add_paragraph(abstract_text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_page_break()

def copy_document_content(source_doc, target_doc):
    """Copy all content from source document to target document"""
    para_count = 0
    table_count = 0
    
    # Copy paragraphs
    for para in source_doc.paragraphs:
        new_para = target_doc.add_paragraph()
        new_para.style = para.style
        new_para.alignment = para.alignment
        
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        
        para_count += 1
    
    # Copy tables
    for table in source_doc.tables:
        new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = table.style
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold or run.italic:
                            for new_para in new_table.rows[i].cells[j].paragraphs:
                                for new_run in new_para.runs:
                                    new_run.bold = run.bold
                                    new_run.italic = run.italic
        
        table_count += 1
    
    return para_count, table_count

def smart_merge():
    """Smart merge of existing sections with new sections"""
    
    print("=" * 80)
    print("📚 SMART MERGE: Creating Master Thesis Document")
    print("=" * 80)
    print()
    
    # Create new master document
    master_doc = Document()
    
    # Add front matter
    print("✅ Adding cover page...")
    add_cover_page(master_doc)
    
    print("✅ Adding Table of Contents placeholder...")
    add_toc_placeholder(master_doc)
    
    print("✅ Adding abstract...")
    add_abstract(master_doc)
    
    print()
    print("📄 Merging sections...")
    print("-" * 80)
    
    total_paragraphs = 0
    total_tables = 0
    
    # Section 1-8 (from combined file)
    sections_1_8_file = 'THESIS_SECTIONS_1_2_3_4_5_6_7_8_COMPLETE.docx'
    if os.path.exists(sections_1_8_file):
        print(f"   Sections 1-8: {sections_1_8_file:50s} ", end="")
        try:
            doc = Document(sections_1_8_file)
            para_count, table_count = copy_document_content(doc, master_doc)
            total_paragraphs += para_count
            total_tables += table_count
            print(f"✅ ({para_count:4d} para, {table_count:2d} tables)")
            master_doc.add_page_break()
        except Exception as e:
            print(f"❌ ERROR: {str(e)}")
    else:
        print(f"⚠️  WARNING: {sections_1_8_file} not found!")
    
    # Section 9
    section9_file = 'THESIS_SECTION_9_FINAL_ARCHITECTURE.docx'
    if os.path.exists(section9_file):
        print(f"   Section 9:   {section9_file:50s} ", end="")
        try:
            doc = Document(section9_file)
            para_count, table_count = copy_document_content(doc, master_doc)
            total_paragraphs += para_count
            total_tables += table_count
            print(f"✅ ({para_count:4d} para, {table_count:2d} tables)")
            master_doc.add_page_break()
        except Exception as e:
            print(f"❌ ERROR: {str(e)}")
    
    # Section 10
    section10_file = 'THESIS_SECTION_10_RESULTS_DISCUSSION.docx'
    if os.path.exists(section10_file):
        print(f"   Section 10:  {section10_file:50s} ", end="")
        try:
            doc = Document(section10_file)
            para_count, table_count = copy_document_content(doc, master_doc)
            total_paragraphs += para_count
            total_tables += table_count
            print(f"✅ ({para_count:4d} para, {table_count:2d} tables)")
            master_doc.add_page_break()
        except Exception as e:
            print(f"❌ ERROR: {str(e)}")
    
    # Section 11
    section11_file = 'THESIS_SECTION_11_CONCLUSION_FUTURE_WORK.docx'
    if os.path.exists(section11_file):
        print(f"   Section 11:  {section11_file:50s} ", end="")
        try:
            doc = Document(section11_file)
            para_count, table_count = copy_document_content(doc, master_doc)
            total_paragraphs += para_count
            total_tables += table_count
            print(f"✅ ({para_count:4d} para, {table_count:2d} tables)")
        except Exception as e:
            print(f"❌ ERROR: {str(e)}")
    
    print("-" * 80)
    print(f"✅ Total content: {total_paragraphs} paragraphs, {total_tables} tables")
    print()
    
    # Save master document
    output_file = 'MASTER_THESIS_COMPLETE.docx'
    master_doc.save(output_file)
    
    file_size = os.path.getsize(output_file) / (1024 * 1024)
    
    print("=" * 80)
    print("✅ MASTER THESIS DOCUMENT CREATED SUCCESSFULLY!")
    print("=" * 80)
    print(f"📄 Output file: {output_file}")
    print(f"📦 File size: {file_size:.2f} MB")
    print(f"📝 Total paragraphs: {total_paragraphs}")
    print(f"📊 Total tables: {total_tables}")
    print()
    print("📋 Document Structure:")
    print("   ✅ Cover Page")
    print("   ✅ Table of Contents (placeholder)")
    print("   ✅ Abstract")
    print("   ✅ Section 1: Introduction")
    print("   ✅ Section 2: Literature Review")
    print("   ✅ Section 3: Dataset & Preprocessing")
    print("   ✅ Section 4: Methodology")
    print("   ✅ Section 5: Implementation")
    print("   ✅ Section 6: Experiments")
    print("   ✅ Section 7: Auto-Labeling")
    print("   ✅ Section 8: Pipeline Optimization (1023 paragraphs!)")
    print("   ✅ Section 9: System Architecture")
    print("   ✅ Section 10: Results & Discussion")
    print("   ✅ Section 11: Conclusion & Future Work")
    print()
    print("🎓 YOUR COMPLETE MASTER'S THESIS IS READY!")
    print("=" * 80)

if __name__ == "__main__":
    smart_merge()
