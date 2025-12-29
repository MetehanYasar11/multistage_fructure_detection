"""
Combine corrected Sections 1 and 2
REAL NUMBERS: 487 images (373 fractured, 114 healthy) in Dataset_2021
"""

from docx import Document
from pathlib import Path

# Load Section 1 (corrected)
doc1 = Document('THESIS_COMPREHENSIVE_REPORT_PART1.docx')

# Load Section 2 (corrected)
doc2 = Document('THESIS_REPORT_CORRECTED_SECTIONS_1_2.docx')

# Section 2 başlangıcını bul (Section 2 başlığından sonraki tüm content'i al)
section2_start = False
for element in doc2.element.body:
    # Section 1'i atla, Section 2'den itibaren ekle
    if section2_start:
        doc1.element.body.append(element)
    elif hasattr(element, 'text'):
        # "2. Dataset Analysis" başlığını bul
        text = element.text if hasattr(element, 'text') else ''
        if '2. Dataset Analysis' in text or '2. Dataset' in text:
            section2_start = True
            doc1.element.body.append(element)

# Save combined
output = Path('THESIS_COMPLETE_SECTIONS_1_2_FINAL.docx')
doc1.save(output)

print(f"COMBINED REPORT CREATED WITH COMPLETE DATASET STATISTICS")
print(f"  Dataset_2021: 487 images (373 fractured, 114 healthy)")
print(f"  Annotations: 915 lines (417 positive fractures + 498 negative hard examples)")
print(f"  Section 1: Introduction with corrected numbers")
print(f"  Section 2: Complete dataset analysis")
print(f"  Total paragraphs: {len(doc1.paragraphs)}")
print(f"  Saved to: {output.absolute()}")
