"""
Embed all visual figures into the master thesis document.
Automatically adds figures with captions and generates List of Figures.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_caption(paragraph, caption_text, figure_number):
    """Add a formatted figure caption"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add "Figure X.Y: " in bold
    run = paragraph.add_run(figure_number)
    run.bold = True
    run.font.size = Pt(10)
    
    # Add caption text in regular
    run = paragraph.add_run(f" {caption_text}")
    run.font.size = Pt(10)
    
    # Add spacing after caption
    paragraph.paragraph_format.space_after = Pt(12)

def find_section_by_heading(doc, heading_text, level=1):
    """Find paragraph index after a specific heading"""
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == f'Heading {level}' and heading_text.lower() in para.text.lower():
            return i + 1  # Return index after heading
    return None

def insert_figure(doc, image_path, width_inches, caption_text, figure_number, insert_after_index):
    """Insert figure with caption at specific position"""
    if not os.path.exists(image_path):
        print(f"⚠️  Warning: Image not found: {image_path}")
        return None
    
    # Insert paragraph at specific position
    new_para = doc.paragraphs[insert_after_index]._element
    new_para_parent = new_para.getparent()
    
    # Create new paragraph for image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    try:
        # Add image
        run = img_para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        
        # Move to correct position
        new_para_parent.insert(new_para_parent.index(new_para) + 1, img_para._element)
        
        # Add caption paragraph
        caption_para = doc.add_paragraph()
        add_caption(caption_para, caption_text, figure_number)
        new_para_parent.insert(new_para_parent.index(img_para._element) + 1, caption_para._element)
        
        print(f"✅ Added: {figure_number} - {caption_text[:50]}...")
        return True
        
    except Exception as e:
        print(f"❌ Error adding figure {figure_number}: {e}")
        return False

def add_list_of_figures(doc):
    """Add List of Figures section after Table of Contents"""
    # Find TOC or add at beginning
    insert_index = 0
    for i, para in enumerate(doc.paragraphs):
        if 'table of contents' in para.text.lower() or 'contents' in para.text.lower():
            insert_index = i + 2  # After TOC and spacing
            break
    
    # Add page break
    if insert_index > 0:
        doc.paragraphs[insert_index].insert_paragraph_before().add_run().add_break()
    
    # Add heading
    heading = doc.add_paragraph()
    heading.style = 'Heading 1'
    heading.text = 'List of Figures'
    
    # Move to correct position
    if insert_index > 0:
        heading_element = heading._element
        doc.paragraphs[insert_index]._element.addprevious(heading_element)
    
    print("✅ Added: List of Figures section")

def embed_all_figures(input_docx, output_docx):
    """Main function to embed all figures"""
    print("="*80)
    print("📊 EMBEDDING FIGURES INTO THESIS DOCUMENT")
    print("="*80)
    
    # Load document
    doc = Document(input_docx)
    print(f"\n📄 Loaded document: {input_docx}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    
    # Define all figures to embed
    figures = [
        # SECTION 1: INTRODUCTION
        {
            'section': 'Introduction',
            'heading': '1.',
            'number': 'Figure 1.1:',
            'path': 'outputs/repo_visualizations/repo_statistics_overview.png',
            'caption': 'Overview of VRF research repository statistics (data sources, experiments, and model variants)',
            'width': 6.0,
            'insert_after': 'clinical significance'
        },
        {
            'section': 'Introduction',
            'heading': '1.',
            'number': 'Figure 1.2:',
            'path': 'outputs/repo_visualizations/research_timeline.png',
            'caption': 'Research and development timeline showing key milestones and model evolution',
            'width': 6.0,
            'insert_after': 'research timeline'
        },
        {
            'section': 'Introduction',
            'heading': '1.',
            'number': 'Figure 1.3:',
            'path': 'outputs/repo_visualizations/experiments_breakdown.png',
            'caption': 'Breakdown of experimental approaches across preprocessing, architecture, and optimization stages',
            'width': 6.0,
            'insert_after': 'experimental'
        },
        
        # SECTION 3: STAGE 1 DETECTION
        {
            'section': 'Stage 1',
            'heading': '3.',
            'number': 'Figure 3.1:',
            'path': 'outputs/debug_detection_first_image.png',
            'caption': 'YOLOv11x root canal treated tooth detection example showing bounding box localization',
            'width': 5.0,
            'insert_after': 'detection performance'
        },
        {
            'section': 'Stage 1',
            'heading': '3.',
            'number': 'Figure 3.2:',
            'path': 'outputs/bbox_scale_analysis.png',
            'caption': 'Bounding box scale analysis demonstrating optimal 2.2× expansion factor for Stage 2 crop extraction',
            'width': 6.0,
            'insert_after': 'scale factor'
        },
        
        # SECTION 4: PREPROCESSING
        {
            'section': 'Preprocessing',
            'heading': '4.',
            'number': 'Figure 4.1:',
            'path': 'outputs/sr_comparison_visualization.png',
            'caption': 'Visual comparison of super-resolution methods (bicubic, ESRGAN, Real-ESRGAN) at 4× upscaling',
            'width': 6.5,
            'insert_after': 'super-resolution'
        },
        {
            'section': 'Preprocessing',
            'heading': '4.',
            'number': 'Figure 4.2:',
            'path': 'outputs/sr_detailed_steps.png',
            'caption': 'Step-by-step visualization of SR+CLAHE preprocessing pipeline showing progressive enhancement',
            'width': 6.5,
            'insert_after': 'pipeline steps'
        },
        {
            'section': 'Preprocessing',
            'heading': '4.',
            'number': 'Figure 4.3:',
            'path': 'outputs/combined_clahe_gabor.png',
            'caption': 'CLAHE+Gabor filter combination showing catastrophic failure with over-enhancement artifacts',
            'width': 6.0,
            'insert_after': 'gabor'
        },
        
        # SECTION 6: EXPERIMENTS & TRAINING
        {
            'section': 'Experiments',
            'heading': '6.',
            'number': 'Figure 6.1:',
            'path': 'runs/vit_classifier/training_history.png',
            'caption': 'ViT-Small training history over 100 epochs showing loss and accuracy convergence',
            'width': 6.0,
            'insert_after': 'training process'
        },
        {
            'section': 'Experiments',
            'heading': '6.',
            'number': 'Figure 6.2:',
            'path': 'runs/vit_classifier/confusion_matrix.png',
            'caption': 'ViT-Small confusion matrix on training test set (231 auto-labeled crops)',
            'width': 5.0,
            'insert_after': 'training test'
        },
        {
            'section': 'Experiments',
            'heading': '6.',
            'number': 'Figure 6.3:',
            'path': 'outputs/risk_zones_vit/stage2_gt_evaluation/stage2_confusion_matrix_gt.png',
            'caption': 'Final validation confusion matrix on 50-image test set showing 84.78% crop-level accuracy (PRIMARY VALIDATION)',
            'width': 5.0,
            'insert_after': 'validation'
        },
        {
            'section': 'Experiments',
            'heading': '6.',
            'number': 'Figure 6.4:',
            'path': 'outputs/risk_zones_vit/stage2_gt_evaluation/stage2_evaluation_summary.png',
            'caption': 'Comprehensive evaluation summary showing per-class metrics and overall performance on 50-image validation',
            'width': 6.0,
            'insert_after': 'evaluation summary'
        },
        
        # SECTION 7: CLASS IMBALANCE
        {
            'section': 'Class Imbalance',
            'heading': '7.',
            'number': 'Figure 7.1:',
            'path': 'runs/class_balancing/class_weights/results.png',
            'caption': 'Class weighting strategy results showing optimal weighted loss [0.73, 1.57] configuration',
            'width': 6.0,
            'insert_after': 'weighted loss'
        },
        {
            'section': 'Class Imbalance',
            'heading': '7.',
            'number': 'Figure 7.2:',
            'path': 'runs/class_balancing/class_weights/confusion_matrix.png',
            'caption': 'Confusion matrix with class weighting demonstrating improved recall (38.9% → 88.71%)',
            'width': 5.0,
            'insert_after': 'recall improvement'
        },
        
        # SECTION 8: PIPELINE OPTIMIZATION
        {
            'section': 'Pipeline Optimization',
            'heading': '8.',
            'number': 'Figure 8.1:',
            'path': 'runs/pipeline_optimization/grid_search_heatmaps.png',
            'caption': 'Grid search heatmaps showing performance across 120 configurations (10 confidence × 12 voting ratios)',
            'width': 6.5,
            'insert_after': 'grid search'
        },
        {
            'section': 'Pipeline Optimization',
            'heading': '8.',
            'number': 'Figure 8.2:',
            'path': 'runs/pipeline_optimization/top_10_configurations.png',
            'caption': 'Top 10 configuration comparisons highlighting optimal balance between sensitivity and specificity',
            'width': 6.0,
            'insert_after': 'top configurations'
        },
        {
            'section': 'Pipeline Optimization',
            'heading': '8.',
            'number': 'Figure 8.3:',
            'path': 'runs/pipeline_optimization/sensitivity_vs_specificity.png',
            'caption': 'Sensitivity vs specificity trade-off analysis showing 8× improvement in specificity (7.69% → 61.54%)',
            'width': 6.0,
            'insert_after': 'specificity improvement'
        },
        {
            'section': 'Pipeline Optimization',
            'heading': '8.',
            'number': 'Figure 8.4:',
            'path': 'runs/pipeline_optimization/confidence_analysis.png',
            'caption': 'Confidence threshold impact analysis on classification performance and decision boundaries',
            'width': 6.0,
            'insert_after': 'confidence threshold'
        },
        {
            'section': 'Pipeline Optimization',
            'heading': '8.',
            'number': 'Figure 8.5:',
            'path': 'runs/pipeline_optimization/optimized_confusion_matrix.png',
            'caption': 'Optimized confusion matrix after combined threshold strategy (conf≥0.75 AND count≥2)',
            'width': 5.0,
            'insert_after': 'optimized'
        },
        
        # SECTION 9: RISK ZONES
        {
            'section': 'Risk Zones',
            'heading': '9.',
            'number': 'Figure 9.1:',
            'path': 'outputs/improved_risk_zones_v2/1069431_1_risk_zones.jpg',
            'caption': 'Risk zone visualization example 1: GREEN zone (Low Risk) - all crops classified as Healthy',
            'width': 5.5,
            'insert_after': 'risk zone'
        },
        {
            'section': 'Risk Zones',
            'heading': '9.',
            'number': 'Figure 9.2:',
            'path': 'outputs/improved_risk_zones_v2/1142690_1_risk_zones.jpg',
            'caption': 'Risk zone visualization example 2: YELLOW zone (Medium Risk) - mixed classifications',
            'width': 5.5,
            'insert_after': 'mixed classification'
        },
        {
            'section': 'Risk Zones',
            'heading': '9.',
            'number': 'Figure 9.3:',
            'path': 'outputs/improved_risk_zones_v2/1159540_risk_zones.jpg',
            'caption': 'Risk zone visualization example 3: RED zone (High Risk) - majority fractured classifications',
            'width': 5.5,
            'insert_after': 'high risk'
        },
        {
            'section': 'Risk Zones',
            'heading': '9.',
            'number': 'Figure 9.4:',
            'path': 'outputs/improved_risk_zones_v2/1260840_1_risk_zones.jpg',
            'caption': 'Risk zone visualization example 4: Clinical decision support with confidence scores',
            'width': 5.5,
            'insert_after': 'confidence'
        },
        
        # SECTION 10: RESULTS
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.1:',
            'path': 'runs/full_pipeline_validation/confusion_matrix.png',
            'caption': 'Full pipeline validation confusion matrix on 50-image test set (crop-level performance)',
            'width': 5.0,
            'insert_after': 'crop-level'
        },
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.2:',
            'path': 'runs/full_pipeline_validation/SCREENING_SYSTEM_ANALYSIS.png',
            'caption': 'Screening system analysis showing image-level aggregation performance (89.47% accuracy)',
            'width': 6.0,
            'insert_after': 'image-level'
        },
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.3:',
            'path': 'runs/full_pipeline_validation/metrics_summary.png',
            'caption': 'Comprehensive metrics summary dashboard comparing crop-level and image-level performance',
            'width': 6.5,
            'insert_after': 'metrics'
        },
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.4:',
            'path': 'outputs/visual_evaluation/000_Fractured_0001.png',
            'caption': 'Qualitative example 1: True positive fractured tooth detection with high confidence',
            'width': 4.5,
            'insert_after': 'qualitative'
        },
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.5:',
            'path': 'outputs/visual_evaluation/001_Fractured_0014.png',
            'caption': 'Qualitative example 2: Fractured tooth with clear visible fracture line',
            'width': 4.5,
            'insert_after': 'fractured'
        },
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.6:',
            'path': 'outputs/visual_evaluation/026_Healthy_0002.png',
            'caption': 'Qualitative example 3: True negative healthy root canal treated tooth',
            'width': 4.5,
            'insert_after': 'healthy'
        },
        {
            'section': 'Results',
            'heading': '10.',
            'number': 'Figure 10.7:',
            'path': 'outputs/visual_evaluation/028_Healthy_0004.png',
            'caption': 'Qualitative example 4: Healthy tooth with no fracture indicators',
            'width': 4.5,
            'insert_after': 'no fracture'
        },
    ]
    
    print(f"\n📊 Total figures to embed: {len(figures)}")
    print("="*80)
    
    # Track statistics
    embedded_count = 0
    failed_count = 0
    missing_count = 0
    
    # Group figures by section for better organization
    current_section = None
    for fig in figures:
        if fig['section'] != current_section:
            current_section = fig['section']
            print(f"\n📍 Section: {current_section}")
            print("-" * 60)
        
        # Check if file exists
        if not os.path.exists(fig['path']):
            print(f"❌ Missing: {fig['number']} - {fig['path']}")
            missing_count += 1
            continue
        
        # Find insertion point
        # Try to find subsection first, then main section
        insert_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.text and fig['insert_after'].lower() in para.text.lower():
                insert_index = i + 1
                break
        
        if insert_index is None:
            # Fallback: find section heading
            for i, para in enumerate(doc.paragraphs):
                if para.style.name.startswith('Heading') and fig['heading'] in para.text:
                    insert_index = i + 5  # Insert after heading and some text
                    break
        
        if insert_index is None:
            print(f"⚠️  Could not find insertion point for {fig['number']}")
            failed_count += 1
            continue
        
        # Insert figure
        result = insert_figure(
            doc, 
            fig['path'], 
            fig['width'], 
            fig['caption'], 
            fig['number'], 
            insert_index
        )
        
        if result:
            embedded_count += 1
        else:
            failed_count += 1
    
    print("\n" + "="*80)
    print(f"📊 EMBEDDING SUMMARY:")
    print(f"   ✅ Successfully embedded: {embedded_count}")
    print(f"   ❌ Failed: {failed_count}")
    print(f"   ⚠️  Missing files: {missing_count}")
    print("="*80)
    
    # Add List of Figures
    print("\n📑 Adding List of Figures section...")
    add_list_of_figures(doc)
    
    # Save document
    print(f"\n💾 Saving document to: {output_docx}")
    doc.save(output_docx)
    print("✅ Document saved successfully!")
    
    print("\n" + "="*80)
    print("🎉 FIGURE EMBEDDING COMPLETE!")
    print("="*80)
    print(f"\n📄 Output: {output_docx}")
    print(f"   Total figures embedded: {embedded_count}/{len(figures)}")
    print("\nNext steps:")
    print("  1. Open the document in Microsoft Word")
    print("  2. Verify figure placements and quality")
    print("  3. Update all fields (Ctrl+A, F9)")
    print("  4. Generate automatic List of Figures (References > Insert Table of Figures)")
    print("  5. Final proofread and export to PDF")
    print("="*80)

if __name__ == "__main__":
    # File paths
    base_dir = Path(__file__).parent
    input_file = base_dir / "thesis_documentation" / "docx" / "MASTER_THESIS_COMPLETE.docx"
    output_file = base_dir / "thesis_documentation" / "docx" / "MASTER_THESIS_WITH_FIGURES.docx"
    
    # Check if input exists
    if not input_file.exists():
        print(f"❌ Error: Input file not found: {input_file}")
        exit(1)
    
    # Run embedding
    embed_all_figures(str(input_file), str(output_file))
