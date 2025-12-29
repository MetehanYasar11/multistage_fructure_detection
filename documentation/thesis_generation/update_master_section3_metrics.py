"""
Update MASTER_THESIS_COMPLETE_ALL_TABLES.docx with real Stage 1 metrics
Replace fake metrics (95%, 98%, 99%) with real metrics (81.05%, 75.77%, 79.06%)
Add training visualization figures
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def update_section3_metrics():
    """Update Section 3 with real Stage 1 detector metrics and visualizations"""
    
    print("🔧 Updating MASTER thesis with real Stage 1 metrics...")
    
    # Load MASTER document
    master_path = Path('docx/MASTER_THESIS_COMPLETE_ALL_TABLES.docx')
    doc = Document(str(master_path))
    
    print(f"📄 Loaded: {master_path}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    
    # Find Section 3 heading
    section3_start = None
    section4_start = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if para.style.name.startswith('Heading 1'):
            if '3.' in text and 'Stage 1' in text:
                section3_start = i
                print(f"✅ Found Section 3 at paragraph {i}: {text[:60]}")
            elif section3_start and ('4.' in text or 'Preprocessing' in text):
                section4_start = i
                print(f"✅ Found Section 4 at paragraph {i}: {text[:60]}")
                break
    
    if not section3_start:
        print("❌ ERROR: Section 3 not found!")
        return
    
    print(f"\n🔍 Section 3 spans paragraphs {section3_start} to {section4_start-1}")
    print(f"   Total paragraphs in Section 3: {section4_start - section3_start}")
    
    # Update metrics in Section 3
    updates_made = 0
    
    for i in range(section3_start, section4_start if section4_start else len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text
        
        # Replace fake metrics with real ones
        replacements = [
            ('0.95 (95%)', '0.8105 (81.05%)'),
            ('95% precision', '81.05% precision'),
            ('0.98 (98%)', '0.7577 (75.77%)'),
            ('98% recall', '75.77% recall'),
            ('0.99 (99%)', '0.7906 (79.06%)'),
            ('99% F1', '79.06% mAP50'),
            ('High Recall (98%)', 'Good Recall (75.77%)'),
            ('High Precision (95%)', 'Strong Precision (81.05%)'),
            ('Near-Perfect F1 (99%)', 'Robust mAP50 (79.06%)'),
            ('only 2% of RCT teeth are missed', 'approximately 24% of RCT teeth are missed'),
            ('only 5% of detections are non-RCT', 'approximately 19% of detections are non-RCT'),
        ]
        
        for old, new in replacements:
            if old in text:
                # Update text
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        updates_made += 1
                        print(f"   ✏️  Updated: '{old}' → '{new}'")
    
    print(f"\n📊 Total metric updates: {updates_made}")
    
    # Find where to add visualizations (after performance analysis)
    insert_position = None
    for i in range(section3_start, section4_start if section4_start else len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if 'Implications for Pipeline' in para.text or 'upper bound on pipeline' in para.text:
            insert_position = i + 2  # After this paragraph and one blank line
            break
    
    if insert_position:
        print(f"\n🖼️  Adding training visualizations at paragraph {insert_position}...")
        
        # Add heading for visualizations
        heading = doc.paragraphs[insert_position].insert_paragraph_before()
        heading.text = '3.9 Training Results and Visualizations'
        heading.style = 'Heading 2'
        
        # Add intro text
        intro = doc.paragraphs[insert_position + 1].insert_paragraph_before()
        intro.text = 'The following figures show the training results from 500 epochs (training_054345):'
        
        # Figure 3.3: Confusion Matrix
        fig_para = doc.paragraphs[insert_position + 2].insert_paragraph_before()
        fig_para.add_run('Figure 3.3: Normalized Confusion Matrix').bold = True
        
        img_path = Path('../outputs/stage1_confusion_matrix_normalized.png')
        if img_path.exists():
            img_para = doc.paragraphs[insert_position + 3].insert_paragraph_before()
            run = img_para.add_run()
            run.add_picture(str(img_path), width=Inches(5.0))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"   ✅ Added: Confusion Matrix")
        
        caption1 = doc.paragraphs[insert_position + 4].insert_paragraph_before()
        caption1.text = "The normalized confusion matrix shows the model's classification performance across all dental structure classes, with strong performance on the RCT class (Class 9)."
        
        # Figure 3.4: Training Curves
        fig_para2 = doc.paragraphs[insert_position + 5].insert_paragraph_before()
        fig_para2.add_run('Figure 3.4: Training Metrics Over 500 Epochs').bold = True
        
        img_path2 = Path('../outputs/stage1_results.png')
        if img_path2.exists():
            img_para2 = doc.paragraphs[insert_position + 6].insert_paragraph_before()
            run2 = img_para2.add_run()
            run2.add_picture(str(img_path2), width=Inches(6.0))
            img_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"   ✅ Added: Training Curves")
        
        caption2 = doc.paragraphs[insert_position + 7].insert_paragraph_before()
        caption2.text = 'Training curves show convergence over 500 epochs. Key observations: (1) box_loss, cls_loss, and dfl_loss steadily decrease; (2) precision, recall, and mAP metrics stabilize around epoch 400-450; (3) final metrics: 81.05% precision, 75.77% recall, 79.06% mAP50.'
        
        # Figure 3.5: F1 Curve
        fig_para3 = doc.paragraphs[insert_position + 8].insert_paragraph_before()
        fig_para3.add_run('Figure 3.5: F1 Score vs Confidence Threshold').bold = True
        
        img_path3 = Path('../outputs/stage1_BoxF1_curve.png')
        if img_path3.exists():
            img_para3 = doc.paragraphs[insert_position + 9].insert_paragraph_before()
            run3 = img_para3.add_run()
            run3.add_picture(str(img_path3), width=Inches(5.0))
            img_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"   ✅ Added: F1-Confidence Curve")
        
        caption3 = doc.paragraphs[insert_position + 10].insert_paragraph_before()
        caption3.text = 'The F1-confidence curve helps select optimal confidence threshold. We use confidence=0.3 in the pipeline to prioritize recall, allowing Stage 2 to filter false positives.'
        
        # Figure 3.6: PR Curve
        fig_para4 = doc.paragraphs[insert_position + 11].insert_paragraph_before()
        fig_para4.add_run('Figure 3.6: Precision-Recall Curve (All Classes)').bold = True
        
        img_path4 = Path('../outputs/stage1_BoxPR_curve.png')
        if img_path4.exists():
            img_para4 = doc.paragraphs[insert_position + 12].insert_paragraph_before()
            run4 = img_para4.add_run()
            run4.add_picture(str(img_path4), width=Inches(5.0))
            img_para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"   ✅ Added: Precision-Recall Curve")
        
        caption4 = doc.paragraphs[insert_position + 13].insert_paragraph_before()
        caption4.text = 'Precision-Recall curves for all dental structure classes. The RCT class (Class 9) shows strong performance with mAP50=79.06%, indicating robust detection across different confidence thresholds and IoU levels.'
        
        print(f"   🎉 Added 4 training visualizations!")
    
    # Save updated document
    output_path = Path('docx/MASTER_THESIS_UPDATED_SECTION3_REAL_METRICS.docx')
    doc.save(str(output_path))
    
    print(f"\n✅ MASTER thesis updated successfully!")
    print(f"   📄 Output: {output_path}")
    print(f"   📊 Real metrics: 81.05% precision, 75.77% recall, 79.06% mAP50")
    print(f"   🖼️  Added 4 training visualizations")
    print(f"\n💡 Original MASTER file unchanged. Review the new file and rename if satisfied.")

if __name__ == '__main__':
    update_section3_metrics()
