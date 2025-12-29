"""
Generate Section 1: Introduction

Simple introduction section that will be merged with others
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_section1():
    doc = Document()
    
    # Setup styles
    styles = doc.styles
    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Calibri'
            style.font.size = Pt(16 - i*2)
            style.font.bold = True
    
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Section 1: Introduction', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 1.1 Background
    doc.add_heading('1.1 Background and Motivation', level=2)
    
    doc.add_paragraph(
        "Dental radiography is a cornerstone of modern dentistry, providing critical diagnostic "
        "information for treatment planning, disease detection, and post-treatment evaluation. "
        "Among various dental pathologies, vertical root fractures (VRF) in root canal treated "
        "teeth present a particularly challenging diagnostic problem due to their subtle appearance "
        "and serious clinical consequences."
    )
    
    doc.add_paragraph(
        "Root canal treatment (RCT) is one of the most common endodontic procedures, performed "
        "to save teeth affected by deep decay or infection. However, RCT-treated teeth are more "
        "susceptible to fractures due to structural weakening from the procedure itself and the "
        "biomechanical changes in dentin properties. Vertical root fractures can lead to treatment "
        "failure, tooth loss, and significant patient morbidity if not detected early."
    )
    
    doc.add_paragraph(
        "Traditional diagnosis of root canal fractures relies on manual inspection of panoramic "
        "dental X-rays by experienced radiologists or endodontists. This process is time-consuming, "
        "subjective, and prone to inter-observer variability. Studies have shown that even experienced "
        "clinicians can miss subtle fracture lines, particularly in early stages when the fracture "
        "gap is minimal. The visual similarity between fractured and healthy RCTs, combined with "
        "image noise and anatomical complexity, makes this a challenging task for human experts."
    )
    
    # 1.2 Problem Statement
    doc.add_heading('1.2 Problem Statement', level=2)
    
    doc.add_paragraph(
        "The clinical challenge of VRF detection can be decomposed into several technical problems:"
    )
    
    problems = [
        "Low Visual Contrast: Fracture lines in panoramic X-rays are often subtle, appearing as "
        "thin dark lines that can be confused with normal anatomical structures or imaging artifacts.",
        
        "Class Imbalance: In typical dental imaging datasets, fractured RCTs are significantly "
        "less common than healthy ones (approximately 30% vs 70%), creating a severe class imbalance "
        "that challenges standard machine learning approaches.",
        
        "Limited Training Data: Medical imaging datasets are inherently small due to privacy concerns, "
        "annotation costs, and the rarity of specific pathologies. Manual labeling of fractures "
        "requires expert knowledge and is extremely time-consuming (40-60 hours for 1,207 crops).",
        
        "Multi-Scale Detection: The diagnostic workflow requires two levels of analysis: "
        "(1) locating individual RCT regions in full panoramic X-rays, and (2) classifying each "
        "RCT as fractured or healthy.",
        
        "Image Quality Variations: Panoramic X-rays suffer from low resolution, poor contrast, and "
        "variable image quality depending on the scanner, acquisition protocol, and patient positioning."
    ]
    
    for i, problem in enumerate(problems, 1):
        doc.add_paragraph(f"{i}. {problem}")
    
    # 1.3 Research Objectives
    doc.add_heading('1.3 Research Objectives', level=2)
    
    doc.add_paragraph(
        "This thesis aims to develop an automated deep learning pipeline for detection and "
        "classification of root canal fractures in panoramic dental X-rays. The specific objectives are:"
    )
    
    objectives = [
        "Develop a two-stage pipeline combining object detection (Stage 1: RCT localization) "
        "and classification (Stage 2: fracture detection)",
        
        "Investigate Vision Transformer architecture for dental fracture classification and "
        "compare with traditional CNN approaches",
        
        "Design and evaluate preprocessing methods to enhance subtle fracture patterns in "
        "low-quality X-ray images",
        
        "Address severe class imbalance through weighted loss functions and systematic comparison "
        "of imbalance handling strategies",
        
        "Create an automated labeling system to accelerate dataset generation and reduce annotation costs",
        
        "Develop a clinical decision support system with risk zone visualization (GREEN/YELLOW/RED) "
        "for intuitive interpretation",
        
        "Conduct comprehensive evaluation on multiple test sets to assess model generalization and "
        "identify deployment considerations"
    ]
    
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f"{i}. {obj}")
    
    # 1.4 Contributions
    doc.add_heading('1.4 Key Contributions', level=2)
    
    doc.add_paragraph(
        "This research makes several novel contributions to dental AI and medical image analysis:"
    )
    
    contributions = [
        "First Application of Vision Transformers to RCT Fracture Detection: Demonstrates that "
        "ViT architecture outperforms CNNs (87.96% vs 85.74% vs 83.72%) for dental fracture classification",
        
        "SR+CLAHE Preprocessing Pipeline: Novel combination of Super-Resolution (4× bicubic) and CLAHE "
        "enhancement achieves +4.63% accuracy improvement over no preprocessing",
        
        "Weighted Loss for Severe Class Imbalance: Systematic comparison of loss functions shows "
        "weighted cross-entropy [0.73, 1.57] improves recall from 38.89% to 88.71% (+50pp)",
        
        "Automated Labeling System (200× Speedup): Liang-Barsky line-clipping algorithm enables "
        "automatic crop labeling in 15 minutes vs 40-60 hours manually, with >95% accuracy",
        
        "Risk Zone Visualization System: Three-tier clinical decision support (GREEN/YELLOW/RED) "
        "provides intuitive output with estimated 30-40% reduction in radiologist review time",
        
        "Comprehensive Evaluation: Multiple test configurations (50-image crop-level, 20-image image-level) "
        "with detailed analysis of Stage 1 detector sensitivity to distribution shift"
    ]
    
    for i, contrib in enumerate(contributions, 1):
        doc.add_paragraph(f"{i}. {contrib}")
    
    # 1.5 Thesis Organization
    doc.add_heading('1.5 Thesis Organization', level=2)
    
    doc.add_paragraph(
        "The remainder of this thesis is organized as follows:"
    )
    
    organization = [
        "Section 2 (Literature Review): Surveys related work in medical image analysis, dental AI, "
        "and VRF detection, positioning this research in the broader context.",
        
        "Section 3 (Dataset & Preprocessing): Documents the Dataset_2021 (487 images, 915 annotations), "
        "manual annotation process (1,207 crops), and data preparation methods.",
        
        "Section 4 (Methodology): Describes the two-stage pipeline architecture, YOLOv11x object detection, "
        "and Vision Transformer classification model.",
        
        "Section 5 (Implementation): Details the PyTorch implementation, training procedures, "
        "hyperparameters, and computational requirements.",
        
        "Section 6 (Experiments): Presents systematic experiments comparing preprocessing methods "
        "(SR+CLAHE vs alternatives), architectures (ViT vs CNNs), and loss functions (weighted vs standard).",
        
        "Section 7 (Auto-Labeling): Introduces the Liang-Barsky automated labeling system, validation "
        "methodology, and dataset expansion results (1,207 → 1,604 crops).",
        
        "Section 8 (Pipeline Optimization): Documents the optimization journey from 7.69% to 61.54% "
        "specificity through grid search (120 combinations) and combined threshold strategies.",
        
        "Section 9 (System Architecture): Presents the complete final system with deployment specifications, "
        "configuration parameters, and integration considerations.",
        
        "Section 10 (Results & Discussion): Analyzes comprehensive evaluation results (84.78% crop-level, "
        "88-94% image-level), Stage 1 detector sensitivity, and clinical implications.",
        
        "Section 11 (Conclusion): Summarizes key findings, contributions, limitations, and outlines "
        "future research directions."
    ]
    
    for org in organization:
        doc.add_paragraph(f"• {org}")
    
    doc.add_paragraph()
    doc.add_paragraph(
        "Through systematic investigation of architectures, preprocessing, and training strategies, "
        "this thesis establishes Vision Transformers as a viable approach for dental fracture detection "
        "and provides a roadmap for clinical deployment of AI-assisted diagnostic tools in dental radiography."
    )
    
    # Save
    output_file = 'THESIS_SECTION_1_INTRODUCTION.docx'
    doc.save(output_file)
    
    print("=" * 80)
    print("✅ SECTION 1 GENERATED SUCCESSFULLY!")
    print("=" * 80)
    print(f"Output file: {output_file}")
    print()
    print("Section 1 Contents:")
    print("  1.1  Background and Motivation")
    print("  1.2  Problem Statement")
    print("  1.3  Research Objectives")
    print("  1.4  Key Contributions")
    print("  1.5  Thesis Organization")
    print("=" * 80)

if __name__ == "__main__":
    generate_section1()
